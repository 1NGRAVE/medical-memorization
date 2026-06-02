"""
生成《医学知识背诵软件 - 完整项目计划书》Word 文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_colored_heading(doc, text, level=1):
    """添加带颜色的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, '2B579A')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacing
    return table

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('医学知识背诵软件')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Medical Memorization App')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('基于 AI 智能判分的医学专业知识背诵系统\n—— 完整项目计划书 ——')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'版本：v1.0\n日期：{datetime.date.today().strftime("%Y年%m月%d日")}\n状态：规划阶段')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
add_colored_heading(doc, '目录', level=1)
toc_items = [
    '一、项目概述',
    '二、核心创新：AI 智能判分系统',
    '三、技术选型与架构',
    '四、项目目录结构',
    '五、数据库设计',
    '六、实施步骤（共 9 步）',
    '七、AI 判分详细设计',
    '八、成本估算',
    '九、UI/UX 关键设计',
    '十、风险与应对',
    '十一、后续迭代方向',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# 一、项目概述
# ============================================================
add_colored_heading(doc, '一、项目概述', level=1)

add_colored_heading(doc, '1.1 项目目标', level=2)
doc.add_paragraph(
    '打造一款专门面向医学生的专业知识背诵软件，类似"百词斩"的交互体验，'
    '但核心差异在于：引入 AI 大语言模型（LLM）进行语义级答案评判，'
    '让医学生不必死记硬背原词原句，而是真正理解知识点后用自己的话表达，'
    'AI 能够灵活判断其回答是否正确并给出详细反馈。'
)
doc.add_paragraph(
    '本项目的另一个核心设计理念是"去中心化AI"：软件本身完全开源免费，'
    '不绑定任何特定 AI 服务商。用户可以根据自己的预算和隐私需求，自由选择'
    '使用免费的云端模型（如 Gemini 免费额度）、付费的商业 API（如 Claude / GPT-4o）、'
    '或完全离线的本地模型（如 Ollama + 开源医学微调模型）。'
    '开发者不承担任何 AI 调用费用，也不经手用户的 API Key 或学习数据。'
)

add_colored_heading(doc, '1.2 解决的问题', level=2)
problems = [
    '传统背诵 APP 只能做字符串匹配，要求逐字逐句背诵，效率低下',
    '医学知识强调理解而非死记，但现有工具无法评判"意思对但措辞不同"的答案',
    '学生缺乏即时、个性化的反馈，不知道自己遗漏了哪些知识点',
    '市面上没有专为医学知识设计的 AI 背诵工具',
]
for p_text in problems:
    doc.add_paragraph(p_text, style='List Bullet')

add_colored_heading(doc, '1.3 核心功能', level=2)
features = [
    '🃏 闪卡模式：传统卡片翻转，手动自评 0-5 分，适配 SM-2 间隔重复算法',
    '🤖 AI 自由作答模式：学生输入文字答案，AI 语义级判分 + 详细反馈',
    '🔌 多模型支持：支持 Claude / GPT / Gemini / DeepSeek 等云端 API，以及 Ollama 本地模型',
    '📝 四选一测验：随机生成干扰项，选错给出 AI 解释',
    '📊 学习统计：每日进度、打卡天数、薄弱知识点分析',
    '📚 多词库管理：解剖学、药理学、病理学、生理学、生物化学',
    '⚙️ 离线降级：无网络或无 API Key 时自动切换到本地关键词匹配模式',
    '🔒 隐私优先：API Key 和学习数据全部存储在用户浏览器本地，不上传任何服务器',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ============================================================
# 二、核心创新：AI 智能判分系统
# ============================================================
add_colored_heading(doc, '二、核心创新：AI 智能判分系统', level=1)

add_colored_heading(doc, '2.1 设计理念', level=2)
doc.add_paragraph(
    '传统背诵工具的核心缺陷是只能做精确字符串匹配，这迫使医学生去背诵教材原文，'
    '而不是理解知识。我们的 AI 判分系统通过大语言模型理解医学概念的语义本质，'
    '能够识别"心肌梗死 = 心脏肌肉因缺血而坏死 = 冠心病导致的急性心肌坏死"等各种表述方式，'
    '真正做到"理解即可得分"。'
)

add_colored_heading(doc, '2.2 三级判分漏斗（推荐方案）', level=2)
doc.add_paragraph(
    '为平衡准确度、延迟和成本，我们采用三级漏斗模型：大部分请求在低成本层级拦截，'
    '只有真正模糊的边界 case 才会调用 LLM。'
)

# 三级漏斗表格
add_styled_table(doc,
    ['层级', '方法', '延迟', '成本/次', '拦截比例', '说明'],
    [
        ['第一关', '关键词硬约束检查', '<1ms', '¥0', '~70%', '全部核心关键词命中→直接通过；全部缺失→直接拒绝'],
        ['第二关', 'Embedding 语义相似度', '<100ms', '≈¥0.0007', '~20%', '相似度>0.85→通过；<0.40→拒绝；中间值→进入第三关'],
        ['第三关', 'LLM 深度评判', '1-3s', '≈¥0.07', '~10%', '调用 Claude/GPT 做最终裁决，返回分数+详细反馈'],
    ]
)

add_colored_heading(doc, '2.3 判分流程图', level=2)
flow_text = """
学生提交答案
    │
    ▼
┌─────────────────────┐
│  第一关：关键词检查   │  ← 本地执行，0延迟
│  提取医学核心关键词    │
└──────┬──────────────┘
       │
    ┌──┴──┐
    ▼     ▼
 全命中   全缺失 ──→ 直接判定(通过/不通过) → 结束
    │     │
    └──┬──┘
      │ 部分命中
      ▼
┌─────────────────────┐
│  第二关：Embedding   │  ← 后端执行，<100ms
│  计算余弦相似度       │
└──────┬──────────────┘
       │
    ┌──┴──┐
    ▼     ▼
  >0.85   <0.40 ──→ 判定(通过/不通过) → 结束
    │     │
    └──┬──┘
      │ 0.40~0.85
      ▼
┌─────────────────────┐
│  第三关：LLM 深度评判 │  ← 调用 Claude API
│  语义分析 + 详细反馈  │  ← 1-3秒，结构化输出
└─────────────────────┘
      │
      ▼
  最终判分 + 反馈 → 结束
"""
p = doc.add_paragraph()
run = p.add_run(flow_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

# ============================================================
# 三、技术选型与架构
# ============================================================
add_colored_heading(doc, '三、技术选型与架构', level=1)

add_colored_heading(doc, '3.1 前端技术栈', level=2)
add_styled_table(doc,
    ['技术', '版本', '用途', '选择理由'],
    [
        ['React', '18+', 'UI 框架', '生态丰富，组件化开发'],
        ['TypeScript', '5+', '类型系统', '减少运行时错误，提升代码可维护性'],
        ['Vite', '5+', '构建工具', '秒级热更新，开发体验极佳'],
        ['Tailwind CSS', '3+', '样式框架', '快速出UI，响应式设计便捷'],
        ['Zustand', '4+', '状态管理', '轻量（<1KB），API简洁，比Redux简单很多'],
        ['Dexie.js', '4+', 'IndexedDB封装', '支持大数据量离线存储，Promise风格API'],
        ['React Router', '6+', '路由管理', '声明式路由，支持嵌套布局'],
        ['Recharts', '2+', '图表库', 'React原生，学习统计可视化'],
        ['Framer Motion', '10+', '动画库', '卡片翻转、页面过渡动画'],
    ]
)

add_colored_heading(doc, '3.2 AI 多模型适配层', level=2)
doc.add_paragraph(
    '本项目不强制依赖任何后端服务器。AI 判分逻辑通过"Provider 适配器模式"实现，'
    '在前端浏览器中直接调用用户配置的 AI 服务。所有 API Key 和判分配置存储在浏览器'
    '的 IndexedDB/localStorage 中，不上传任何服务器。'
)
add_styled_table(doc,
    ['Provider 类型', '具体模型', '费用', '延迟', '需要 API Key', '适用场景'],
    [
        ['云端付费', 'Claude 3.5 Haiku / GPT-4o-mini / DeepSeek V3', '用户自付', '1-3s', '是', '追求最佳判分质量'],
        ['云端免费', 'Gemini 2.5 Flash（免费额度）', '🆓 免费', '1-3s', '是（免费注册）', '零成本入门'],
        ['本地模型', 'Ollama + Llama 3 / Qwen 3 / 开源医学模型', '🆓 完全免费', '<500ms（本地）', '否', '隐私优先 / 离线使用'],
        ['浏览器内置', 'Chrome 内置 Gemini Nano（实验性）', '🆓 免费', '<100ms', '否', '未来零延迟方案'],
    ]
)
doc.add_paragraph(
    'Provider 适配器统一接口：所有 Provider 实现相同的 judge(studentAnswer, referenceAnswer) => JudgeResult '
    '方法，上层业务代码无需关心底层用的是哪个模型。切换模型只需在设置页修改配置，无需改代码。'
)

add_colored_heading(doc, '3.3 纯前端架构（零后端）', level=2)

add_colored_heading(doc, '3.3 架构图', level=2)
arch_text = """
                    ┌─────────────────────────────────────────┐
                    │          纯前端 React SPA（静态托管）       │
                    │  GitHub Pages / Vercel / Netlify 免费部署  │
                    └─────────────────────────────────────────┘
                                      │
            ┌─────────────────────────┼─────────────────────────┐
            ▼                         ▼                         ▼
   ┌─────────────────┐     ┌─────────────────┐     ┌─────────────────┐
   │   📚 数据层       │     │   🧠 AI 判分层   │     │   📊 UI 层       │
   │                  │     │                  │     │                  │
   │ Dexie.js         │     │ Provider适配器    │     │ 首页/背诵/统计    │
   │ (IndexedDB)      │     │                  │     │                  │
   │ · decks          │     │ ┌──────────────┐ │     │ React Router     │
   │ · cards          │     │ │ 关键词检查    │ │     │ Framer Motion    │
   │ · review_logs    │     │ │ (本地, 0ms)  │ │     │ Recharts         │
   │ · settings       │     │ └──────┬───────┘ │     │                  │
   │                  │     │        ▼         │     │                  │
   │ SM-2 算法引擎     │     │ ┌──────────────┐ │     │                  │
   │                  │     │ │ Embedding    │ │     │                  │
   │                  │     │ │ (可选, 100ms)│ │     │                  │
   │                  │     │ └──────┬───────┘ │     │                  │
   │                  │     │        ▼         │     │                  │
   │                  │     │ ┌──────────────┐ │     │                  │
   │                  │     │ │ LLM Provider │ │     │                  │
   │                  │     │ │ (用户自选)    │ │     │                  │
   │                  │     │ └──────────────┘ │     │                  │
   └─────────────────┘     └────────┬────────┘     └─────────────────┘
                                    │  用户浏览器直接调用 AI API
                       ┌────────────┼────────────┬──────────────┐
                       ▼            ▼            ▼              ▼
                 ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐
                 │ Claude   │ │ GPT-4o   │ │ Gemini   │ │ Ollama   │
                 │ API      │ │ Mini     │ │ (免费)    │ │ (本地)   │
                 │ 💰付费   │ │ 💰付费   │ │ 🆓免费   │ │ 🆓离线   │
                 └──────────┘ └──────────┘ └──────────┘ └──────────┘
                        ▲            ▲            ▲              ▲
                        │            │            │              │
                        └────────────┴────────────┴──────────────┘
                              用户自行配置 API Key（存于浏览器本地）
"""
p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# ============================================================
# 四、项目目录结构
# ============================================================
add_colored_heading(doc, '四、项目目录结构', level=1)

dir_tree = """medical-memorization/
├── index.html                      # HTML 入口
├── package.json                    # 项目配置与依赖
├── tsconfig.json                   # TypeScript 配置
├── vite.config.ts                  # Vite 构建配置
├── tailwind.config.js              # Tailwind CSS 配置
│
├── src/                            # 🖥️ React 前端（纯静态，零后端）
│   ├── main.tsx                    # 应用入口
│   ├── App.tsx                     # 根组件 + 路由配置
│   ├── routes/
│   │   ├── HomePage.tsx            # 首页 — 今日学习概览
│   │   ├── StudyPage.tsx           # 闪卡背诵模式
│   │   ├── FreeAnswerPage.tsx      # 🤖 AI 自由作答模式（核心）
│   │   ├── QuizPage.tsx            # 四选一测验模式
│   │   ├── StatsPage.tsx           # 学习统计与可视化
│   │   └── SettingsPage.tsx        # 设置（模型选择 / API Key / 目标）
│   ├── components/
│   │   ├── layout/
│   │   │   ├── BottomNav.tsx       # 底部导航栏
│   │   │   └── PageHeader.tsx      # 页面顶栏
│   │   ├── study/
│   │   │   ├── FlashCard.tsx       # 闪卡组件（3D翻转动画）
│   │   │   ├── ScoreButtons.tsx    # 评分按钮组（0-5分）
│   │   │   └── ProgressBar.tsx     # 进度条
│   │   ├── ai-judge/
│   │   │   ├── AnswerInput.tsx     # 自由文本输入组件
│   │   │   ├── JudgeResult.tsx     # AI 判分结果展示
│   │   │   ├── ModelSelector.tsx   # 🔌 模型选择器组件
│   │   │   ├── KeywordHighlighter.tsx  # 关键词高亮
│   │   │   └── JudgeSkeleton.tsx   # 判分加载骨架屏
│   │   ├── quiz/
│   │   │   ├── QuizCard.tsx        # 题目卡片
│   │   │   ├── OptionButton.tsx    # 选项按钮
│   │   │   └── QuizResult.tsx      # 测验结果
│   │   └── common/
│   │       ├── Modal.tsx           # 通用弹窗
│   │       └── Toast.tsx           # 消息提示
│   ├── stores/
│   │   ├── studyStore.ts           # 学习状态管理
│   │   ├── deckStore.ts            # 词库状态管理
│   │   └── settingsStore.ts        # 设置状态管理（含模型配置）
│   ├── providers/                  # 🔌 AI 多模型适配层（核心设计）
│   │   ├── types.ts                # Provider 统一接口定义
│   │   ├── registry.ts             # Provider 注册表
│   │   ├── keyword-check.ts        # 关键词硬约束检查（模型无关，始终运行）
│   │   ├── claude-provider.ts      # Claude API Provider
│   │   ├── openai-provider.ts      # OpenAI API Provider
│   │   ├── gemini-provider.ts      # Gemini API Provider（含免费额度）
│   │   ├── deepseek-provider.ts    # DeepSeek API Provider
│   │   ├── ollama-provider.ts      # Ollama 本地模型 Provider
│   │   └── judge-pipeline.ts       # 三级判分流水线（关键词→Embedding→LLM）
│   ├── db/
│   │   ├── index.ts                # Dexie 数据库实例
│   │   ├── deckService.ts          # 词库 CRUD
│   │   ├── cardService.ts          # 卡片 CRUD
│   │   └── reviewService.ts        # 复习记录
│   ├── algorithms/
│   │   └── sm2.ts                  # SM-2 间隔重复算法
│   ├── data/
│   │   └── decks/                  # 内置医学词库 JSON
│   │       ├── anatomy.json        # 解剖学
│   │       ├── pharmacology.json   # 药理学
│   │       ├── pathology.json      # 病理学
│   │       ├── physiology.json     # 生理学
│   │       └── biochemistry.json   # 生物化学
│   └── types/
│       └── index.ts                # TypeScript 类型定义
│
└── docs/                           # 📄 文档
    └── plan.md                     # 本计划书 Markdown 版本
"""
p = doc.add_paragraph()
run = p.add_run(dir_tree)
run.font.name = 'Consolas'
run.font.size = Pt(8)

doc.add_page_break()

# ============================================================
# 五、数据库设计
# ============================================================
add_colored_heading(doc, '五、数据库设计（IndexedDB）', level=1)

add_colored_heading(doc, '5.1 Decks 表（词库）', level=2)
add_styled_table(doc,
    ['字段', '类型', '说明'],
    [
        ['id', 'string (UUID)', '主键'],
        ['name', 'string', '词库名称，如"药理学"'],
        ['description', 'string', '词库描述'],
        ['icon', 'string', '图标 emoji 或 URL'],
        ['cardCount', 'number', '卡片总数'],
        ['color', 'string', '主题色'],
        ['createdAt', 'Date', '创建时间'],
    ]
)

add_colored_heading(doc, '5.2 Cards 表（卡片）', level=2)
add_styled_table(doc,
    ['字段', '类型', '说明'],
    [
        ['id', 'string (UUID)', '主键'],
        ['deckId', 'string', '所属词库 ID'],
        ['question', 'string', '问题/名词，如"简述普萘洛尔的作用机制"'],
        ['referenceAnswer', 'string', '标准答案（完整版）'],
        ['keyPoints', 'string[]', '关键知识点列表，如["β受体阻断","膜稳定","支气管禁忌"]'],
        ['keywords', 'string[]', '必须包含的核心关键词'],
        ['answerMode', 'enum', '"flashcard" | "free_answer" | "both"'],
        ['imageUrl', 'string?', '配图 URL（解剖图等）'],
        ['difficulty', 'number', '难度系数 (1-5)'],
        ['source', 'string?', '知识来源/教材引用'],
        # SM-2 字段
        ['repetitions', 'number', '连续正确次数'],
        ['interval', 'number', '当前间隔（天）'],
        ['ease', 'number', '难度系数（默认2.5）'],
        ['nextReview', 'Date', '下次复习日期'],
        ['lastReview', 'Date?', '上次复习日期'],
    ]
)

add_colored_heading(doc, '5.3 ReviewLogs 表（复习记录）', level=2)
add_styled_table(doc,
    ['字段', '类型', '说明'],
    [
        ['id', 'string (UUID)', '主键'],
        ['cardId', 'string', '卡片 ID'],
        ['mode', 'enum', '"flashcard" | "free_answer" | "quiz"'],
        ['score', 'number', '得分 (0-5)'],
        ['studentAnswer', 'string?', '学生作答内容（AI模式）'],
        ['aiFeedback', 'string?', 'AI 反馈内容'],
        ['aiMissedPoints', 'string[]?', 'AI 指出的遗漏点'],
        ['aiCorrections', 'json?', 'AI 指出的错误纠正'],
        ['previousInterval', 'number', '调整前间隔'],
        ['newInterval', 'number', '调整后间隔'],
        ['reviewedAt', 'Date', '复习时间'],
        ['responseTime', 'number?', '回答耗时（秒）'],
    ]
)

doc.add_page_break()

# ============================================================
# 六、实施步骤
# ============================================================
add_colored_heading(doc, '六、实施步骤（共 9 步）', level=1)

steps = [
    {
        'title': '第 1 步：项目初始化与基础架构',
        'time': '预计 1-2 天',
        'tasks': [
            '使用 Vite 创建 React + TypeScript 项目',
            '配置 Tailwind CSS、ESLint、Prettier',
            '安装核心依赖：Zustand、Dexie.js、React Router、Recharts、Framer Motion',
            '搭建 Express 后端骨架（server/ 目录）',
            '配置 Vite 代理转发 /api → Express',
            '搭建路由骨架（5 个页面占位）',
            '定义核心 TypeScript 类型（Card, Deck, ReviewLog, JudgeResult 等）',
            '创建 .env.example 模板',
        ]
    },
    {
        'title': '第 2 步：数据库设计与 SM-2 算法',
        'time': '预计 1-2 天',
        'tasks': [
            '实现 Dexie 数据库初始化与表结构迁移',
            '编写 deckService / cardService / reviewService CRUD 操作',
            '实现 SM-2 间隔重复算法核心逻辑',
            '编写 SM-2 单元测试（验证间隔与 ease 调整正确性）',
            '实现"获取今日到期卡片"查询函数',
        ]
    },
    {
        'title': '第 3 步：多模型 AI 判分层（⭐核心）',
        'time': '预计 2-3 天',
        'tasks': [
            '定义 JudgeProvider 统一接口（src/providers/types.ts）',
            '实现 keyword-check.ts：提取标准答案关键词，检查学生答案覆盖率（模型无关，始终运行）',
            '实现 claude-provider.ts：封装 Claude API 调用，结构化 JSON 输出',
            '实现 openai-provider.ts：封装 OpenAI API 调用',
            '实现 gemini-provider.ts：封装 Gemini API 调用（免费额度）',
            '实现 ollama-provider.ts：封装 Ollama 本地模型调用',
            '实现 registry.ts：Provider 注册表，根据设置动态激活',
            '实现 judge-pipeline.ts：三级判分流水线（关键词→可选Embedding→LLM Provider）',
            '编写 System Prompt（医学教授角色 + 5条评判原则 + JSON 输出格式）',
            '实现 LRU 判分缓存（相同答案不重复调用 API）',
            '实现 Provider 连接测试功能（设置页"测试连接"按钮）',
        ]
    },
    {
        'title': '第 4 步：首页 — 今日学习概览',
        'time': '预计 1 天',
        'tasks': [
            '环形进度条组件（今日已学 / 今日目标）',
            '连续打卡天数显示',
            '"开始背诵"大按钮（带动画）',
            '各词库学习进度条',
            '底部导航栏（背诵 / 测验 / 统计 / 设置）',
            '今日 AI 判分次数与剩余次数显示',
        ]
    },
    {
        'title': '第 5 步：闪卡背诵模式',
        'time': '预计 1-2 天',
        'tasks': [
            '3D 卡片翻转动画（CSS perspective + rotateY）',
            '正面显示问题，背面显示标准答案 + 配图',
            '底部 4 个评分按钮：「忘记了」「模糊」「记得」「太简单」',
            '评分后调用 SM-2 算法更新卡片，自动切换到下一张',
            '今日任务完成时显示庆祝动画（撒花效果）',
            '中途退出时保存进度',
        ]
    },
    {
        'title': '第 6 步：AI 自由作答模式（⭐核心页面）',
        'time': '预计 2-3 天',
        'tasks': [
            '题目展示区：显示问题 + 难度 + 所属词库',
            '自由文本输入框（支持 Markdown 快捷输入）',
            '"提交评判"按钮 → 显示骨架屏加载 → 展示判分结果',
            '判分结果卡片：分数动画 + 详细反馈 + 遗漏点列表 + 错误纠正',
            '关键词高亮对比（学生答案 vs 标准答案关键词覆盖情况）',
            '离线降级提示条（无网络时自动切换到关键词模式）',
            '"显示提示"按钮（可选：展示部分关键词但不展示完整答案）',
            '历史作答记录查看（同一张卡片的历次AI评分变化）',
            '响应式适配（移动端可用）',
        ]
    },
    {
        'title': '第 7 步：四选一测验模式',
        'time': '预计 1-2 天',
        'tasks': [
            '随机抽取 10 张到期卡片作为一轮测验',
            '干扰项生成算法（同词库随机选 3 张其他卡片的答案）',
            '选对 → 绿色动画 + 0.5s 后进入下一题',
            '选错 → 红色标记 + 显示正确选项 + 调用 AI 生成简短解释',
            '一轮结束显示得分 + 错题回顾',
            '错题自动加入今日复习队列',
        ]
    },
    {
        'title': '第 8 步：学习统计页',
        'time': '预计 1-2 天',
        'tasks': [
            '折线图：最近 7/30 天每日学习卡片数（Recharts）',
            '热力图：每日学习时长（类似 GitHub 贡献图）',
            '饼图：各词库掌握比例',
            'AI 判分准确率趋势折线图',
            '薄弱知识点 Top 5（AI 最常指出的遗漏点）',
            '累计数据卡片：总卡片数、已掌握数、学习总时长、AI 判分次数',
            '学习建议（基于 AI 错误模式的简单规则生成）',
        ]
    },
    {
        'title': '第 9 步：医学词库 + 设置页 + 部署',
        'time': '预计 2-3 天',
        'tasks': [
            '编写 4-5 个内置词库 JSON（每库 20-50 张卡片）',
            '每张卡片包含：question、referenceAnswer、keyPoints、keywords',
            '设置页：AI 模型选择（下拉框 + "测试连接"按钮）',
            '设置页：API Key 配置（按 Provider 分别配置，加密存储在 localStorage）',
            '设置页：每日目标、主题切换（浅/深）、判分严格度调节',
            '数据导入导出（JSON 格式，方便教师分发词库，用户迁移数据）',
            '用量仪表盘（本月 API 调用次数、预估费用）',
            '重置学习进度确认弹窗',
            '整体 UI 打磨、动画优化',
            'README.md 编写（含部署说明 + 各 AI 模型配置教程）',
            '部署到 GitHub Pages / Vercel（纯静态，零配置）',
        ]
    },
]

for i, step in enumerate(steps):
    add_colored_heading(doc, f'{i+1}. {step["title"]}', level=2)
    p = doc.add_paragraph()
    run = p.add_run(f'⏱️ {step["time"]}')
    run.bold = True
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    for task in step['tasks']:
        doc.add_paragraph(task, style='List Bullet')

doc.add_page_break()

# ============================================================
# 七、AI 判分详细设计
# ============================================================
add_colored_heading(doc, '七、AI 判分详细设计', level=1)

add_colored_heading(doc, '7.1 System Prompt 设计', level=2)
doc.add_paragraph('以下是 LLM 判分使用的 System Prompt（医学教授角色）：')

prompt_text = """你是一位经验丰富的医学院教授，正在使用语音/文字考核学生的专业知识。
你的任务是判断学生的回答是否在语义层面与标准答案一致。

## 评判原则（按优先级排列）

1. 【语义优先】只判断含义是否一致，不要求措辞相同。
   示例："心肌梗死" = "心脏肌肉缺血坏死" = "冠心病导致的心肌急性坏死" ✅

2. 【关键点拆分】标准答案由多个关键知识点组成（见key_points字段），
   逐一检查学生是否覆盖，每遗漏一个关键点扣1分。

3. 【部分给分】答对一部分但遗漏其他部分，按覆盖比例给分。
   覆盖 ≥80% → 4分，60-79% → 3分，40-59% → 2分，<40% → 1-0分。

4. 【容错合理表述】学生用自己的话正确表达，即使术语不够专业也要认可。
   "心脏跳动的起始点"可以等同于"窦房结"。

5. 【致命错误零容忍】概念性错误必须明确指出并扣分。
   把"动脉"说成"静脉"、"激动剂"说成"抑制剂"，即使其他都对也要严厉扣分。

## 输出格式（严格JSON，不要其他内容）
{
  "score": 整数0-5,
  "is_pass": 布尔值(≥3分为true),
  "coverage_rate": 0.0-1.0的关键点覆盖率,
  "feedback": "2-3句话的鼓励性评语，先肯定优点再指出不足",
  "missed_points": ["遗漏的关键知识点"],
  "corrections": [
    {
      "student_said": "学生原文",
      "should_be": "正确表述",
      "note": "简短解释"
    }
  ]
}"""

p = doc.add_paragraph()
run = p.add_run(prompt_text)
run.font.name = 'Consolas'
run.font.size = Pt(8)

add_colored_heading(doc, '7.2 关键词检查算法', level=2)
doc.add_paragraph('第一关关键词检查的核心逻辑（本地执行，0延迟）：')

kw_text = """function keywordCheck(studentAnswer: string, keywords: string[]): {
  pass: boolean,      // 是否直接通过
  fail: boolean,      // 是否直接失败
  coverage: number,   // 命中率 0.0-1.0
  matched: string[],  // 命中的关键词
  missed: string[]    // 未命中的关键词
} {
  const lower = studentAnswer.toLowerCase()
  const matched = []
  const missed = []

  for (const kw of keywords) {
    // 使用同义词库扩展匹配
    const synonyms = getSynonyms(kw)  // 如"心肌梗死" → ["心梗","MI","myocardial infarction"]
    const found = synonyms.some(s => lower.includes(s.toLowerCase()))
    if (found) {
      matched.push(kw)
    } else {
      missed.push(kw)
    }
  }

  const coverage = matched.length / keywords.length

  return {
    pass: coverage >= 1.0,       // 全部命中 → 直接通过
    fail: coverage <= 0.1,       // 几乎全没命中 → 直接失败
    coverage,
    matched,
    missed
  }
}"""

p = doc.add_paragraph()
run = p.add_run(kw_text)
run.font.name = 'Consolas'
run.font.size = Pt(8)

add_colored_heading(doc, '7.3 Embedding 相似度判断', level=2)
doc.add_paragraph(
    '使用文本向量模型（如 OpenAI text-embedding-3-small 或本地模型）将学生答案和标准答案'
    '分别转换为向量，然后计算余弦相似度。阈值设定：'
)
doc.add_paragraph('相似度 ≥ 0.85 → 判定正确（语义高度一致）', style='List Bullet')
doc.add_paragraph('相似度 < 0.40 → 判定错误（语义差异显著）', style='List Bullet')
doc.add_paragraph('0.40 ≤ 相似度 < 0.85 → 边界情况，送交 LLM 最终裁决', style='List Bullet')

doc.add_page_break()

# ============================================================
# 八、成本估算
# ============================================================
add_colored_heading(doc, '八、开发者与用户成本分析', level=1)

add_colored_heading(doc, '8.1 开发者成本：≈ ¥0', level=2)
doc.add_paragraph(
    '本软件为纯前端静态应用，无需任何后端服务器。可免费部署在：'
)
deploy_options = [
    'GitHub Pages：完全免费，无限流量',
    'Vercel：免费套餐（100GB 带宽/月），足够个人和小团队使用',
    'Netlify：免费套餐（100GB 带宽/月）',
    'Cloudflare Pages：完全免费，无限流量',
    '任何静态文件托管服务',
]
for opt in deploy_options:
    doc.add_paragraph(opt, style='List Bullet')
doc.add_paragraph(
    '开发者唯一可能的成本是域名费用（可选，约 ¥50/年），以及如果未来想提供官方云服务'
    '（非必需，用户自带 API Key 即可）。'
)

add_colored_heading(doc, '8.2 用户成本：完全自控', level=2)
doc.add_paragraph(
    '用户根据自己的选择和用量付费，开发者不承担任何 AI 调用成本：'
)

add_styled_table(doc,
    ['用户选择', '每日50题成本', '月成本（30天）', '年成本', '是否需要 API Key'],
    [
        ['🆓 Gemini 2.5 Flash（免费额度）', '¥0', '¥0', '¥0', '需免费注册 Google 账号'],
        ['🆓 Ollama 本地模型', '¥0（仅电费）', '¥0', '¥0', '不需要'],
        ['💰 DeepSeek V3（极低价）', '≈¥0.005', '≈¥0.15', '≈¥1.8', '需要，充值 ¥10 可用数年'],
        ['💰 GPT-4o Mini', '≈¥0.05', '≈¥1.5', '≈¥18', '需要 OpenAI API Key'],
        ['💰 Claude Haiku', '≈¥0.05', '≈¥1.5', '≈¥18', '需要 Anthropic API Key'],
    ]
)

doc.add_paragraph(
    '结论：用户可以零成本使用（Gemini 免费额度 或 Ollama 本地模型），即使选择最贵的商业 API，'
    '一个月也只要 ¥1.5-18，不到一杯奶茶的钱。开发者完全零负担。'
)

add_colored_heading(doc, '8.3 成本控制功能（内置）', level=2)
strategies = [
    '📊 用量仪表盘：显示本月 API 调用次数和预估费用',
    '🔔 费用预警：用户可设置月度预算上限，超出后自动切换本地判分',
    '💾 判分缓存：相同/相似答案直接复用，减少重复 API 调用（预计节省 40-60%）',
    '⚡ 关键词预筛：70% 的答案通过关键词检查直接判定，无需调用 AI',
    '🔌 随时切换：任何时刻都可在设置中更换模型或切换至离线模式',
]
for s in strategies:
    doc.add_paragraph(s, style='List Bullet')

# Add a new section: Multi-Provider Architecture
doc.add_page_break()

add_colored_heading(doc, '八-附录：多模型适配器架构详解', level=1)

add_colored_heading(doc, '8A.1 Provider 统一接口设计', level=2)
doc.add_paragraph(
    '所有 AI Provider 实现相同的 TypeScript 接口，上层业务代码无需关心底层使用的是哪个模型。'
    '切换模型只需修改设置页的配置，无需改动任何业务代码。'
)

provider_interface = """// src/providers/types.ts — 统一 Provider 接口

interface JudgeProvider {
  /** Provider 唯一标识 */
  id: string;
  /** 显示名称（设置页下拉框中显示） */
  name: string;
  /** 是否需要 API Key */
  requiresApiKey: boolean;
  /** 是否需要网络连接 */
  requiresNetwork: boolean;
  /** 费用标签 */
  costLabel: 'free' | 'cheap' | 'paid';

  /** 核心方法：评判学生答案 */
  judge(params: {
    studentAnswer: string;
    referenceAnswer: string;
    keyPoints: string[];
    keywords: string[];
  }): Promise<JudgeResult>;

  /** 测试连接是否可用（设置页"测试连接"按钮） */
  testConnection(apiKey?: string): Promise<boolean>;
}"""

p = doc.add_paragraph()
run = p.add_run(provider_interface)
run.font.name = 'Consolas'
run.font.size = Pt(7.5)

add_colored_heading(doc, '8A.2 Provider 注册表', level=2)
doc.add_paragraph('内置 Provider 列表，用户可在设置页一键切换：')

add_styled_table(doc,
    ['Provider', '模型', '费用', '需要Key', '离线可用', '备注'],
    [
        ['claude', 'Claude 3.5 Haiku', '💰 付费', '是', '否', '判分质量最佳'],
        ['openai', 'GPT-4o Mini', '💰 付费', '是', '否', '性价比高'],
        ['gemini', 'Gemini 2.5 Flash', '🆓 免费', '是（免费注册）', '否', '零成本首选'],
        ['deepseek', 'DeepSeek V3', '💰 极低价', '是', '否', '中文医学知识优秀'],
        ['ollama', '用户本地模型', '🆓 免费', '否', '是', '隐私最佳，完全离线'],
    ]
)

add_colored_heading(doc, '8A.3 判分流水线（Provider无关）', level=2)
doc.add_paragraph(
    '无论用户选择哪个 Provider，判分流水线的第一关（关键词检查）始终在本地运行，'
    '只有需要深度评判的边界 case 才会调用用户配置的 AI Provider：'
)

pipeline_text = """// src/providers/judge-pipeline.ts

export async function judgeAnswer(
  studentAnswer: string,
  card: Card,
  provider: JudgeProvider  // 用户选择的 Provider，可任意替换
): Promise<JudgeResult> {

  // 第一关：关键词硬约束检查（始终在本地运行，0ms，0成本）
  const kwResult = keywordCheck(studentAnswer, card.keywords)
  if (kwResult.pass) return { score: 5, feedback: '完美!', ... }
  if (kwResult.fail) return { score: 0, feedback: '请重新复习...', ... }

  // 第二关：如果 Provider 支持 Embedding，先算语义相似度（可选）
  if (provider.supportsEmbedding) {
    const similarity = await provider.embed(studentAnswer, card.referenceAnswer)
    if (similarity > 0.85) return { score: 4, feedback: '很好!', ... }
    if (similarity < 0.40) return { score: 1, feedback: '需要加强...', ... }
  }

  // 第三关：调用用户配置的 AI Provider 深度评判
  return await provider.judge({
    studentAnswer,
    referenceAnswer: card.referenceAnswer,
    keyPoints: card.keyPoints,
    keywords: card.keywords,
  })
}"""

p = doc.add_paragraph()
run = p.add_run(pipeline_text)
run.font.name = 'Consolas'
run.font.size = Pt(7)

doc.add_page_break()

# ============================================================
# 九、UI/UX 关键设计
# ============================================================
add_colored_heading(doc, '九、UI/UX 关键设计', level=1)

add_colored_heading(doc, '9.1 AI 自由作答页面原型', level=2)
wireframe = """
┌──────────────────────────────────────────┐
│  ← 返回          药理学 · 第 3/20 张      │
├──────────────────────────────────────────┤
│                                          │
│  ┌────────────────────────────────────┐  │
│  │  📝 请简述普萘洛尔（Propranolol）  │  │
│  │  的作用机制和临床应用。             │  │
│  │                                    │  │
│  │  💡 提示：请从受体类型、药理作用、  │  │
│  │  临床用途三个方面回答。            │  │
│  └────────────────────────────────────┘  │
│                                          │
│  ┌────────────────────────────────────┐  │
│  │                                    │  │
│  │  (学生在此区域自由输入答案)         │  │
│  │                                    │  │
│  │  · 支持 Markdown 格式              │  │
│  │  · 支持快捷键提交 (Ctrl+Enter)     │  │
│  │                                    │  │
│  │                                    │  │
│  └────────────────────────────────────┘  │
│                                          │
│  [💡 显示提示]              [📤 提交评判]│
│                                          │
│  ┌─ 🤖 AI 评判结果 ──────────────────┐  │
│  │                                    │  │
│  │  ⭐ 得分: 4 / 5                    │  │
│  │  📊 关键点覆盖率: 75%              │  │
│  │                                    │  │
│  │  ✨ 你的回答正确描述了 β 受体阻断  │  │
│  │  机制和心血管临床应用。需要注意：  │  │
│  │                                    │  │
│  │  📝 遗漏的知识点:                  │  │
│  │  · 膜稳定作用 (MSA)               │  │
│  │  · 支气管哮喘患者禁用             │  │
│  │                                    │  │
│  │  🔧 概念纠正:                     │  │
│  │  你说"选择性阻断β1受体"→           │  │
│  │  应为"非选择性阻断β1和β2受体"      │  │
│  │                                    │  │
│  └────────────────────────────────────┘  │
│                                          │
│                  [📖 下一题 →]           │
└──────────────────────────────────────────┘"""

p = doc.add_paragraph()
run = p.add_run(wireframe)
run.font.name = 'Consolas'
run.font.size = Pt(7.5)

add_colored_heading(doc, '9.2 关键交互细节', level=2)
interactions = [
    '3D 卡片翻转：使用 CSS perspective + rotateY，翻转角度 180°，过渡 0.6s',
    'AI 判分加载：骨架屏 + 逐字打字动画（模拟 AI 正在思考）',
    '评分后自动下一张：300ms 后自动切换，减少操作步骤',
    '离线降级提示：页面顶部黄色提示条，非阻塞式',
    '庆祝动画：今日目标达成时撒花/烟花粒子效果（canvas-confetti 库）',
    '移动端适配：所有页面响应式设计，最小支持 375px 宽度',
]
for inter in interactions:
    doc.add_paragraph(inter, style='List Bullet')

doc.add_page_break()

# ============================================================
# 十、风险与应对
# ============================================================
add_colored_heading(doc, '十、风险与应对', level=1)

add_styled_table(doc,
    ['风险', '影响', '概率', '应对策略'],
    [
        ['LLM 幻觉（给错误答案打高分）', '高', '中',
         '关键词硬约束兜底：核心概念缺失 → 强制 ≤2分；支持用户申诉反馈不合理的判分结果'],
        ['API 延迟不稳定（>5秒）', '中', '低',
         '设置超时自动降级到本地关键词判分；前端显示"正在深度分析…"加载动画；Ollama 本地模型无此问题'],
        ['不同 AI 模型判分标准不一致', '中', '中',
         '关键词检查层作为统一基准线；System Prompt 标准化确保各模型遵循相同规则；用户可切换模型对比判分结果'],
        ['用户不知如何获取 API Key', '低', '高',
         '设置页内置图文教程链接；README 提供各平台注册指南；Ollama 本地模型零配置作为兜底方案'],
        ['Gemini 免费额度未来可能取消', '低', '低',
         '多 Provider 架构天然不受单一服务商影响；随时可切换到其他免费/低价 Provider'],
        ['医学数据隐私（发送到云端 API）', '高', '低',
         '所有数据存储在浏览器本地；Ollama Provider 完全离线判分；仅发送答题文本，不包含用户个人信息'],
        ['医学术语 LLM 理解不准', '中', '中',
         'Prompt 中加入中英文医学术语对照表；关键词检查层包含同义词库；提供学生申诉/反馈通道持续优化'],
        ['冷启动无用户数据', '低', '高',
         '内置高质量医学词库（含 reference_answer 和 key_points）；后续可建立社区词库共享机制'],
    ]
)

doc.add_page_break()

# ============================================================
# 十一、后续迭代方向
# ============================================================
add_colored_heading(doc, '十一、后续迭代方向', level=1)

add_colored_heading(doc, '11.1 V1.0 → V2.0', level=2)
v1_features = [
    '🎙️ 语音输入作答：语音转文字后提交 AI 判分，模拟口试场景',
    '📱 移动端 App：使用 Capacitor 打包为 iOS/Android 原生应用',
    '☁️ 多设备同步：添加后端数据库（PostgreSQL）+ 用户账号系统',
    '🤝 协作词库：教师创建/分享词库，学生订阅学习',
    '🧠 AI 智能出题：根据卡片内容，AI 自动生成多种题型（简答/填空/病例分析）',
    '📋 错题本：自动收集错误卡片，智能安排复习频率',
    '🏆 排行榜/成就系统：学习动力激励',
]
for f in v1_features:
    doc.add_paragraph(f, style='List Bullet')

add_colored_heading(doc, '11.2 远期愿景', level=2)
vision = [
    'AI 虚拟导师：基于学生历史错误模式，生成个性化学习路径',
    '病例闯关模式：将知识点串联成临床病例，模拟诊断推理过程',
    '知识图谱：可视化展示医学概念之间的关联关系',
    '多语言支持：中文/英文/拉丁文医学名词对照学习',
    '开源社区：允许医生和教师贡献词库与判分标准',
]
for v in vision:
    doc.add_paragraph(v, style='List Bullet')

# ============================================================
# 页脚
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 文档结束 —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(10)

# ============================================================
# 保存
# ============================================================
output_path = 'd:/Project/26.6.1medical memorization/医学知识背诵软件_完整项目计划书.docx'
doc.save(output_path)
import sys
sys.stdout.reconfigure(encoding='utf-8')
print(f'Word document generated: {output_path}')
